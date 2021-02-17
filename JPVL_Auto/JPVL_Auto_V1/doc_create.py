from f_input import collect_user_input, title, input_type
from docx import Document
from docx.shared import Pt
from file_processing import original_words, kana_and_eng_def


def create_vocab_list(): 
    name = title[0] #This is the title and file name which was set in f_input. It was nessesary to append the string to a list
    document = Document()
    document.add_heading(name, 0) #Sets heading of document to value set in title


    #Make a tuple of tuples
    word_list = tuple(zip(original_words, kana_and_eng_def))
    print(word_list)

    #Create table
    table = document.add_table(rows=1, cols=2) #Row number will increase automatically as content is looped in
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Word' #This and below line set the title for each of the two columns
    hdr_cells[1].text = 'Reading + Meanings'
    hdr_cells[0].width = 1097280 #This and below line define how wide the columns are 
    hdr_cells[1].width = 48463200 
    
    for word, meaning in word_list: 
        #Here we loop through the tuple of tuples of word and meanings and add them to the table
        row_cells = table.add_row().cells
        row_cells[0].text = word
        row_cells[1].text = meaning


    for row in table.rows: 
        #This changes the font size of the table to 9 (reduces the number of pages)
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(9)
    
    #Below we set the file name of the vocabulary list (if title box empty, default file name is 'vocabularly list)
    output_file_extension = '.docx'
    file_name = name + output_file_extension
    if not name:
        document.save('Vocabulary list.docx')
    else:
        document.save(file_name)
